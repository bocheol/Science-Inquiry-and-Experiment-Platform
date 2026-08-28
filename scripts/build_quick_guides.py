from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "manuals"
SERVICE_URL = "https://science-inquiry-platform-974188506094.asia-northeast3.run.app"

GREEN = "146B55"
GREEN_DARK = "0E4F40"
GREEN_SOFT = "EAF5F0"
BLUE_SOFT = "EDF3FA"
GOLD_SOFT = "FFF6D8"
RED_SOFT = "FCECEB"
INK = "1E2D28"
MUTED = "5F7069"
LINE = "CBD8D2"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=INK, name="Malgun Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("쪽 " )
    set_run_font(run, size=8.5, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_base = max(existing_abs or [0]) + 10
    num_base = max(existing_num or [0]) + 10

    def create_definition(abstract_id, num_id, fmt, text, left, hanging, color):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(tabs)
        p_pr.append(ind)
        r_pr = OxmlElement("w:rPr")
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color)
        bold = OxmlElement("w:b")
        r_pr.append(color_el)
        r_pr.append(bold)
        for node in (start, num_fmt, lvl_text, suff, p_pr, r_pr):
            lvl.append(node)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create_definition(abstract_base, num_base, "decimal", "%1.", 430, 260, GREEN)
    create_definition(abstract_base + 1, num_base + 1, "bullet", "•", 400, 220, GREEN)
    return num_base, num_base + 1


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def make_document(audience):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in (
        ("Title", 24, GREEN_DARK, 0, 5),
        ("Subtitle", 11.5, MUTED, 0, 12),
        ("Heading 1", 15, GREEN, 11, 6),
        ("Heading 2", 11.5, GREEN_DARK, 7, 3),
    ):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"과탐실 AI 탐구 플랫폼  |  {audience} 빠른 안내")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)
    num_step, num_bullet = configure_numbering(doc)
    return doc, num_step, num_bullet


def add_title_block(doc, audience, subtitle):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("2026학년도 통합과학 팀 탐구")
    set_run_font(run, size=9, bold=True, color=GREEN)
    title = doc.add_paragraph(style="Title")
    title.add_run(f"{audience} 빠른 안내서")
    sub = doc.add_paragraph(style="Subtitle")
    sub.add_run(subtitle)
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [10140], indent_dxa=120)
    set_table_borders(callout, color="B9D9CB", size=8)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, GREEN_SOFT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("접속 주소  ")
    set_run_font(run, size=10, bold=True, color=GREEN_DARK)
    add_hyperlink(p, SERVICE_URL, SERVICE_URL)
    p2 = cell.add_paragraph("태블릿·노트북의 Chrome 또는 기본 브라우저에서 접속하세요.")
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        set_run_font(run, size=8.8, color=MUTED)


def add_step(doc, num_id, title, intro=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title)
    set_run_font(run, size=11.5, bold=True, color=GREEN_DARK)
    if intro:
        run = p.add_run(f"  {intro}")
        set_run_font(run, size=9.5, color=MUTED)


def add_bullet(doc, num_id, text, bold_prefix=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = Cm(0.72)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, size=9.8, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, size=9.8)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.8)
    return p


def add_note(doc, title, body, fill=GOLD_SOFT, border="E6CE72"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [10140], indent_dxa=120)
    set_table_borders(table, color=border, size=7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=9.5, bold=True, color=GREEN_DARK if fill != RED_SOFT else "8A3030")
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    for run in p2.runs:
        set_run_font(run, size=9.2, color=INK)


def add_status_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2400, 7740], indent_dxa=120)
    set_table_borders(table)
    for index, heading in enumerate(("화면 상태", "뜻과 다음 행동")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE_SOFT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(heading)
        set_run_font(run, size=9.2, bold=True, color=GREEN_DARK)
    for label, detail in rows:
        cells = table.add_row().cells
        for index, text in enumerate((label, detail)):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            set_run_font(run, size=8.8, bold=index == 0)
    set_table_geometry(table, [2400, 7740], indent_dxa=120)
    return table


def add_checklist_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)


def build_teacher_guide(path):
    doc, steps, bullets = make_document("교사용")
    add_title_block(doc, "교사용", "팀 편성 → 학생 시작 → 탐구 계획서 검토를 수업 전에 빠르게 확인합니다.")
    add_checklist_heading(doc, "수업 전 5분 점검")
    add_bullet(doc, bullets, "교사 아이디로 로그인하고 대시보드가 열리는지 확인합니다.")
    add_bullet(doc, bullets, "학생 로그인 카드를 준비하고, 태블릿이 학교 WiFi에 연결되는지 확인합니다.")
    add_bullet(doc, bullets, "학급별 팀과 팀장이 맞는지 확인한 뒤 학생에게 접속 주소를 안내합니다.")

    add_step(doc, steps, "교사 로그인", "처음 받은 임시 비밀번호는 반드시 변경합니다.")
    add_bullet(doc, bullets, "로그인 화면에서 교사 아이디와 비밀번호를 입력합니다.")
    add_bullet(doc, bullets, "최초 로그인이라면 8자 이상, 글자와 숫자가 포함된 새 비밀번호로 바꿉니다.")
    add_bullet(doc, bullets, "교사 계정은 4명 모두 1~9반을 함께 관리할 수 있습니다.")

    add_step(doc, steps, "팀 편성 및 팀장 지정")
    add_bullet(doc, bullets, "교사 대시보드 → ‘팀 편성’에서 학급을 선택합니다.")
    add_bullet(doc, bullets, "필요하면 ‘+ 팀 추가’를 누르고, 학생 행의 ‘현재 팀’ 목록에서 팀을 배정합니다.")
    add_bullet(doc, bullets, "팀원 한 명을 ‘팀장’으로 지정합니다. 기본은 4인 1팀이며 인원 변동도 가능합니다.")
    add_bullet(doc, bullets, "학생 이동·제거 기록은 삭제되지 않고 교사용 이력으로 남습니다.")
    add_note(doc, "주의", "운영 화면에서 시험용 팀을 만들거나 실제 학생 자료를 삭제하지 마세요. 팀을 잘못 배정했다면 해당 학생의 ‘현재 팀’을 올바른 팀으로 다시 선택합니다.")

    doc.add_page_break()
    add_checklist_heading(doc, "계획서 진행 확인과 승인")
    add_step(doc, steps, "학생 진행 상황 확인")
    add_bullet(doc, bullets, "대시보드의 ‘학급별 탐구 진척’에서 학급·팀, 탐구 주제, 계획서 상태를 확인합니다.")
    add_bullet(doc, bullets, "팀 행의 ‘팀 확인’을 누르면 AI 대화, 계획서, 준비물 신청을 함께 볼 수 있습니다.")
    add_status_table(doc, [
        ("작성 중", "학생 팀이 계획서를 작성하고 있습니다."),
        ("승인 대기", "학생이 제출했습니다. 내용을 검토해 주세요."),
        ("수정 요청", "교사 피드백에 따라 학생 팀이 수정합니다."),
        ("재승인 필요", "승인 뒤 학생이 내용을 바꾸어 다시 검토해야 합니다."),
        ("승인됨", "계획이 확정되어 학생 개인 실험 일지가 열립니다."),
    ])

    add_step(doc, steps, "계획서 검토")
    add_bullet(doc, bullets, "팀 상세의 ‘탐구 계획서’에서 주제, 동기·목적, 이론적 배경과 출처, 방법·변인, 일정, 기대효과, 참고문헌을 확인합니다.")
    add_bullet(doc, bullets, "바로 진행해도 되면 ‘계획서 승인’을, 보완이 필요하면 구체적인 피드백을 적고 ‘수정 요청’을 누릅니다.")
    add_bullet(doc, bullets, "승인 뒤 상태를 바꿀 때는 ‘승인 상태 변경’을 누르고 화면에 표시된 ‘반 + 팀명’을 정확히 입력해야 합니다.")

    add_step(doc, steps, "수업 중 자주 생기는 문제")
    add_bullet(doc, bullets, "학생 비밀번호 분실: 팀 편성 표의 ‘비밀번호 초기화’로 새 임시 비밀번호를 발급합니다.")
    add_bullet(doc, bullets, "학생이 팀 화면을 못 봄: 현재 팀 배정과 계정 상태를 먼저 확인합니다.")
    add_bullet(doc, bullets, "AI 답변 생성 중 입력이 안 됨: 정상적인 순차 처리입니다. 답변이 끝날 때까지 기다립니다.")
    add_bullet(doc, bullets, "저장·전송 오류: 학생이 작성 내용을 지우지 않게 하고, 연결 복구 후 다시 시도합니다.")

    add_note(doc, "개인정보 원칙", "학생 이름·학번은 필요한 교사 화면과 DB에서만 사용합니다. AI에는 팀원 가명만 전달됩니다. 개인 일지는 작성자와 교사만 볼 수 있고, 연락처·이메일은 계획서에서 수집하지 않습니다.", fill=GREEN_SOFT, border="B9D9CB")
    doc.save(path)


def build_student_guide(path):
    doc, steps, bullets = make_document("학생용")
    add_title_block(doc, "학생용", "로그인 → 팀 주제 찾기 → 공동 계획서 작성과 제출 순서로 진행합니다.")
    add_step(doc, steps, "첫 로그인과 비밀번호 변경")
    add_bullet(doc, bullets, "선생님에게 받은 로그인 카드의 5자리 학번과 임시 비밀번호를 입력합니다.")
    add_bullet(doc, bullets, "처음 로그인하면 새 비밀번호를 만듭니다: 8자 이상, 글자와 숫자를 함께 사용합니다.")
    add_bullet(doc, bullets, "비밀번호는 친구와 공유하지 않습니다. 잊어버리면 선생님에게 초기화를 요청합니다.")

    add_step(doc, steps, "내 팀 확인")
    add_bullet(doc, bullets, "로그인 후 화면 위쪽에서 반·팀명·팀원과 팀장을 확인합니다.")
    add_bullet(doc, bullets, "팀이 다르거나 ‘아직 팀이 배정되지 않았습니다’가 나오면 선생님께 바로 알립니다.")

    add_step(doc, steps, "AI와 탐구 방향 찾기")
    add_bullet(doc, bullets, "‘이론 탐구’ 탭에서 팀이 궁금해하는 현상이나 관심 분야를 팀의 말로 적습니다.")
    add_bullet(doc, bullets, "‘AI와 방향 3개 찾기’를 누르고, 제안된 연구 질문·변인·안전 내용을 팀원과 비교합니다.")
    add_bullet(doc, bullets, "팀이 합의한 카드에서 ‘이 방향 선택’을 누른 뒤 AI와 질문을 이어 갑니다.")
    add_bullet(doc, bullets, "한 번에 한 질문만 처리됩니다. AI 답변이 생성되는 동안에는 새 입력이 잠시 잠깁니다.")
    add_note(doc, "AI 사용 원칙", "AI는 완성된 계획서를 대신 써 주는 사람이 아닙니다. 우리 팀의 생각을 먼저 적고, 모르는 개념·변인·오차·안전·출처를 질문하여 계획을 더 구체적으로 만드세요.", fill=GREEN_SOFT, border="B9D9CB")

    doc.add_page_break()
    add_checklist_heading(doc, "팀 탐구 계획서 작성")
    add_step(doc, steps, "항목별 공동 작성")
    add_bullet(doc, bullets, "‘탐구 계획’ 탭을 열고 팀원이 역할을 나누어 항목별로 작성합니다.")
    add_bullet(doc, bullets, "항목을 편집한 뒤 다른 곳을 누르면 저장됩니다. 같은 항목을 다른 팀원이 편집 중이면 작성자 표시와 잠금이 나타납니다.")
    add_bullet(doc, bullets, "팀명·팀원·지도교사 정보는 자동으로 연결됩니다. 연락처와 이메일은 입력하지 않습니다.")

    add_step(doc, steps, "계획서에 꼭 들어갈 내용")
    add_bullet(doc, bullets, "왜 탐구하는지: 연구 동기와 목적")
    add_bullet(doc, bullets, "무엇을 알고 시작하는지: 이론적 배경·선행 연구·실제 출처")
    add_bullet(doc, bullets, "어떻게 확인할지: 측정할 변인, 바꿀 조건, 같게 유지할 조건, 자료 수집 방법")
    add_bullet(doc, bullets, "언제 무엇을 할지: 일정·장소·탐구 내용·준비물")
    add_bullet(doc, bullets, "예상 결과·기대효과와 참고문헌")

    add_step(doc, steps, "제출과 교사 피드백")
    add_bullet(doc, bullets, "팀원과 모든 항목을 확인한 뒤 ‘선생님께 제출’을 누릅니다.")
    add_status_table(doc, [
        ("선생님 확인 중", "제출이 완료되었습니다. 검토 결과를 기다립니다."),
        ("수정 요청", "교사 피드백을 읽고 해당 항목을 보완한 뒤 다시 제출합니다."),
        ("승인 완료", "계획이 승인되었습니다. 이후 실험 일지 작성이 열립니다."),
        ("재승인 필요", "승인 뒤 계획을 수정했습니다. 반드시 다시 제출합니다."),
    ])

    add_step(doc, steps, "제출 전 마지막 확인")
    add_bullet(doc, bullets, "연구 질문이 측정하거나 비교할 수 있을 만큼 구체적인가?")
    add_bullet(doc, bullets, "실제 확인한 출처만 적었고, 존재하지 않는 논문·사이트를 만들지 않았는가?")
    add_bullet(doc, bullets, "위험한 화학물질·불꽃·고전압·미생물·인체 적용이 있다면 교사 확인을 받았는가?")
    add_bullet(doc, bullets, "실명·학번·연락처 같은 개인정보를 AI 대화에 직접 입력하지 않았는가?")
    add_note(doc, "연결이 불안할 때", "오류가 보여도 작성 내용을 먼저 지우지 마세요. WiFi를 확인하고 잠시 뒤 다시 저장하거나 제출합니다. 반복해서 버튼을 누르기보다 화면의 저장 상태와 오류 문구를 선생님께 보여 주세요.")
    doc.save(path)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_teacher_guide(OUTPUT_DIR / "2026-과탐실-교사용-빠른-안내서.docx")
    build_student_guide(OUTPUT_DIR / "2026-과탐실-학생용-빠른-안내서.docx")
    print(OUTPUT_DIR)


if __name__ == "__main__":
    main()
