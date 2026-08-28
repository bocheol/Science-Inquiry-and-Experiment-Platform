from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "manuals"
SHOTS = OUT / "screenshots"
SERVICE_URL = "https://science-inquiry-platform-974188506094.asia-northeast3.run.app"

# compact_reference_guide with named Korean-school overrides:
# A4 paper, Malgun Gothic, green platform palette, 17.6 cm content width.
GREEN = "146B55"
GREEN_DARK = "0E4F40"
GREEN_SOFT = "EAF5F0"
BLUE_SOFT = "EDF3FA"
GOLD_SOFT = "FFF6D8"
RED_SOFT = "FCECEB"
INK = "1E2D28"
MUTED = "5F7069"
LINE = "CBD8D2"
WHITE = "FFFFFF"
CONTENT_DXA = 9980
TABLE_INDENT = 120


def set_run_font(run, size=None, bold=None, color=INK, italic=None, name="Malgun Gothic"):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("쪽 ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_base = max(abs_ids or [0]) + 20
    num_base = max(num_ids or [0]) + 20

    def definition(abstract_id, num_id, fmt, text):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend((tabs, ind, spacing))
        r_pr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), GREEN)
        r_pr.append(color)
        level.extend((start, num_fmt, lvl_text, suff, p_pr, r_pr))
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)

    definition(abstract_base, num_base, "bullet", "•")
    definition(abstract_base + 1, num_base + 1, "decimal", "%1.")
    return num_base, num_base + 1


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def base_document(audience):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 28, GREEN_DARK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 16, GREEN, 18, 10),
        ("Heading 2", 13, GREEN_DARK, 14, 7),
        ("Heading 3", 11.5, GREEN_DARK, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run(f"과탐실 AI 탐구 플랫폼  |  {audience} 상세 사용서"), size=8.3, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])
    doc.core_properties.title = f"과탐실 AI 탐구 플랫폼 {audience} 상세 사용서"
    doc.core_properties.subject = "2026학년도 통합과학 팀 탐구 플랫폼 사용 안내"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    bullets, steps = configure_numbering(doc)
    return doc, bullets, steps


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_cover(doc, audience, subtitle, chapters):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(62)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(kicker.add_run("2026학년도 통합과학 팀 탐구"), size=10, bold=True, color=GREEN)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(f"{audience} 상세 사용서")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(subtitle)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run("과탐실 AI 탐구 플랫폼 · 운영 리비전 2026-08-29"), size=9.5, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [CONTENT_DXA])
    set_table_borders(callout, color="B9D9CB", size=7)
    cell = callout.cell(0, 0)
    shade_cell(cell, GREEN_SOFT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("접속 주소  "), size=10, bold=True, color=GREEN_DARK)
    add_hyperlink(p, "과탐실 AI 탐구 플랫폼 열기", SERVICE_URL)
    p2 = cell.add_paragraph("Chrome 또는 태블릿의 기본 브라우저에서 접속합니다. 비밀번호와 학생 자료는 사용서에 기록하지 않습니다.")
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        set_run_font(run, size=9.1, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(toc.add_run("이 사용서에서 다루는 내용"), size=11.5, bold=True, color=GREEN_DARK)
    for chapter in chapters:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(chapter), size=9.5, color=INK)
    doc.add_page_break()


def add_chapter(doc, title, intro, *, page_break=True):
    if page_break and len(doc.paragraphs) > 1 and doc.paragraphs[-1].text:
        doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(title)
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        set_run_font(run, size=10.4, color=MUTED)


def add_step(doc, steps, title, text):
    p = doc.add_paragraph()
    apply_numbering(p, steps)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(title), size=11.2, bold=True, color=GREEN_DARK)
    if text:
        set_run_font(p.add_run(f"  {text}"), size=9.6, color=MUTED)


def add_bullets(doc, bullets, items):
    for text in items:
        p = doc.add_paragraph()
        apply_numbering(p, bullets)
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(text), size=10.0)


def add_note(doc, title, body, kind="info"):
    fill, border, title_color = {
        "info": (GREEN_SOFT, "B9D9CB", GREEN_DARK),
        "warning": (GOLD_SOFT, "E6CE72", "765918"),
        "danger": (RED_SOFT, "E6A7A3", "8A3030"),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=border, size=7)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title), size=9.7, bold=True, color=title_color)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        set_run_font(run, size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_screenshot(doc, filename, caption, alt_text, width_cm=17.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(SHOTS / filename), width=Cm(width_cm))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(8)
    set_run_font(cp.add_run(caption), size=8.7, italic=True, color=MUTED)


def add_status_table(doc, rows, headings=("화면 상태", "뜻과 다음 행동")):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2500, CONTENT_DXA - 2500])
    set_table_borders(table)
    repeat_header(table.rows[0])
    for i, heading in enumerate(headings):
        cell = table.rows[0].cells[i]
        shade_cell(cell, BLUE_SOFT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(heading), size=9.2, bold=True, color=GREEN_DARK)
    for label, detail in rows:
        cells = table.add_row().cells
        for i, text in enumerate((label, detail)):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), size=9.0, bold=i == 0)
    set_table_geometry(table, [2500, CONTENT_DXA - 2500])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_checklist(doc, bullets, title, items):
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(title)
    add_bullets(doc, bullets, [f"□ {item}" for item in items])


def build_teacher(path):
    chapters = [
        "1. 로그인·계정과 팀 편성", "2. AI 탐구·계획서·준비물", "3. 일지·보고서·변경 이력",
        "4. 진척 대시보드·시험", "5. 자기·동료평가", "6. 도움말·문제 해결·보안",
    ]
    doc, bullets, steps = base_document("교사용")
    add_cover(doc, "교사용", "수업 준비부터 평가 결과 공개까지 전체 학기 흐름을 운영하는 방법", chapters)

    add_chapter(doc, "1. 로그인·계정과 팀 편성", "교사 4명은 1~9반을 공동 관리합니다. 학생이 수업을 시작하기 전에 계정·학급·팀장 배정을 확인합니다.")
    add_step(doc, steps, "교사 로그인", "현재 사용하는 교사 아이디와 본인 비밀번호를 입력합니다.")
    add_bullets(doc, bullets, [
        "최초 로그인이라면 8자 이상이며 글자와 숫자가 포함된 새 비밀번호로 변경합니다.",
        "교사 비밀번호는 다른 교사와 공유하지 않고, 브라우저 공용 저장 기능에 남기지 않습니다.",
        "학생이 비밀번호를 잊은 경우에만 해당 학생 행에서 임시 비밀번호를 재발급합니다.",
    ])
    add_screenshot(doc, "00-login.png", "그림 1. 운영 로그인 화면", "왼쪽에 플랫폼 소개, 오른쪽에 아이디와 비밀번호 입력란이 있는 운영 로그인 화면")
    add_step(doc, steps, "팀 만들기·배정·팀장 지정", "학급을 선택한 뒤 실제 수업 편성대로 배정합니다.")
    add_bullets(doc, bullets, [
        "팀을 만들고 학생 행의 현재 팀을 선택합니다. 기본은 4인 1팀이지만 수업 상황에 맞게 조정할 수 있습니다.",
        "현재 팀원 중 한 명을 팀장으로 지정합니다. 계획서·보고서 복원 권한은 현재 팀장에게만 있습니다.",
        "학생을 이동·제거해도 과거 대화·일지·역할·평가 기록은 교사용으로 보존됩니다.",
        "운영 DB나 Google Sheet에 시험용 팀·학생 자료를 만들지 않습니다. 체험은 별도 체험 계정을 사용합니다.",
    ])
    add_note(doc, "수업 전 확인", "학생 화면 상단에 올바른 반·팀·팀원·팀장이 보이는지 확인합니다. 잘못 배정했다면 학생 자료를 삭제하지 말고 현재 팀만 올바르게 다시 선택합니다.", "warning")

    add_chapter(doc, "2. AI 탐구·계획서·준비물", "학생 팀은 관심사를 입력해 탐구 방향 3개를 받고, 선택한 방향을 바탕으로 계획서를 공동 작성합니다.")
    add_step(doc, steps, "AI 이론 탐구 지도", "AI는 질문과 힌트를 제공하고 완성 계획서를 대신 쓰지 않습니다.")
    add_bullets(doc, bullets, [
        "팀 관심사에 연구 질문·변인·안전 요소가 드러나도록 구체적인 문장을 적게 합니다.",
        "AI가 제시한 출처는 학생이 원문을 열어 실제 내용과 작성 주체를 확인하게 합니다.",
        "학생 실명·학번·연락처를 AI 대화에 입력하지 않도록 안내합니다. 플랫폼은 팀원을 가명으로 전달합니다.",
        "AI 응답 생성 중에는 다음 입력이 잠시 잠깁니다. 답변이 끝난 뒤 이어서 질문합니다.",
    ])
    add_screenshot(doc, "01-student-overview.png", "그림 2. 익명 예시로 본 학생 이론 탐구 화면", "익명 체험팀의 탐구 탭과 팀 공유 AI 대화, 선택한 탐구 방향이 표시된 화면")
    add_step(doc, steps, "계획서 검토·승인", "주제·동기·이론·변인·방법·일정·안전·출처를 확인합니다.")
    add_bullets(doc, bullets, [
        "보완이 필요하면 무엇을 어떻게 고칠지 구체적으로 적고 수정 요청합니다.",
        "승인 뒤 학생이 내용을 바꾸면 재승인 필요 상태로 바뀝니다.",
        "승인 상태를 되돌릴 때는 화면에 표시된 반과 팀명을 정확히 입력해야 합니다.",
    ])
    add_status_table(doc, [
        ("작성 중", "학생 팀이 아직 작성 중입니다."), ("승인 대기", "학생이 제출했습니다. 교사가 검토합니다."),
        ("수정 요청", "피드백에 따라 학생이 보완합니다."), ("재승인 필요", "승인 뒤 내용이 바뀌어 다시 확인해야 합니다."),
        ("승인", "실험 일지와 보고서 단계로 진행할 수 있습니다."),
    ])
    add_step(doc, steps, "준비물 신청 확인", "품목과 예산, 링크, Google Sheet 전송 상태를 확인합니다.")
    add_bullets(doc, bullets, [
        "모바일 쇼핑 앱의 공유 문구는 지원 쇼핑몰이면 PC 상품 주소로 자동 정리됩니다.",
        "변환할 수 없는 모바일 전용 주소는 학생에게 PC 브라우저에서 상품 주소를 다시 복사하도록 안내합니다.",
        "조별 합계가 5만원을 넘으면 경고가 표시됩니다. 필요성과 대체품을 확인합니다.",
        "Google Sheet 전송 실패는 교사 화면에서 재전송합니다. 반복 전송 전 기존 행을 먼저 확인합니다.",
    ])

    add_chapter(doc, "3. 일지·보고서·변경 이력", "계획 승인 뒤 학생은 개인 일지를 작성하고, 팀은 학교 양식 순서대로 최종보고서를 공동 작성합니다.")
    add_step(doc, steps, "개인 실험 일지", "일지는 작성자 학생과 교사만 볼 수 있습니다.")
    add_bullets(doc, bullets, [
        "차시·날짜·오늘 한 일·관찰 결과·느낀 점과 궁금한 점을 학생 개인이 기록합니다.",
        "차시당 사진은 최대 5장입니다. 태블릿의 큰 사진은 업로드 전에 자동으로 줄어듭니다.",
        "연결이 잠시 끊기면 작성 내용이 브라우저에 보존되고, 연결 복구 뒤 다시 전송됩니다.",
        "제거된 학생은 이후 팀 자료에 접근할 수 없지만 과거 일지는 교사가 계속 확인할 수 있습니다.",
    ])
    add_step(doc, steps, "팀 최종보고서", "연구 목적부터 부록까지 항목별로 나누어 공동 작성합니다.")
    add_bullets(doc, bullets, [
        "같은 항목은 한 번에 한 명이 편집하고, 다른 항목은 여러 팀원이 동시에 작성할 수 있습니다.",
        "팀원별 역할을 확인한 뒤 제출합니다. 교사는 확인 완료 또는 수정 요청을 선택합니다.",
        "수정 요청 뒤 보완하면 학생 팀이 다시 제출해야 합니다.",
    ])
    add_step(doc, steps, "변경 이력과 복원", "계획서와 보고서는 저장 직전의 전체 상태를 이력으로 남깁니다.")
    add_screenshot(doc, "03-document-history.png", "그림 3. 계획서 변경 이력과 복원 버튼", "익명 예시의 계획서 변경 이력 세 건과 이 상태로 복원 버튼")
    add_bullets(doc, bullets, [
        "복원은 교사와 해당 팀의 현재 활성 팀장만 할 수 있습니다. 일반 팀원과 이전 팀장은 복원할 수 없습니다.",
        "복원 직전 상태도 다시 이력으로 남으므로 연속 복원이 가능합니다.",
        "복원 뒤에는 승인·확인 상태가 초기화되며 다시 제출·검토합니다.",
    ])

    add_chapter(doc, "4. 진척 대시보드·시험", "대시보드에서 교사 처리 항목과 학생 진행 항목을 분리해 보고, 필요한 범위만 내보냅니다.")
    add_screenshot(doc, "02-teacher-dashboard.png", "그림 4. 학급별 탐구 진척 대시보드", "익명 팀 두 개의 계획 승인, 일지 작성, 보고서 제출과 교사 확인 상태를 보여 주는 대시보드")
    add_bullets(doc, bullets, [
        "학급과 상태 필터로 승인 대기·재승인·보고서 검토·준비물 오류·예산 초과를 확인합니다.",
        "Excel에는 팀 진척과 학생 일지 현황 시트가, CSV에는 팀 진척 요약이 들어갑니다.",
        "내보내기에는 일지 본문·사진·비밀번호가 포함되지 않습니다.",
    ])
    add_step(doc, steps, "시험 생성", "공통·팀 공통·개인화 문항 수와 총점을 정해 초안을 만듭니다.")
    add_bullets(doc, bullets, [
        "전체 공통 문항은 특정 팀 자료를 그대로 사용하지 않고 중립 자료를 바탕으로 생성합니다.",
        "팀 공통 문항은 해당 팀의 승인 계획서와 보고서, 개인화 문항은 해당 학생 본인의 일지와 역할만 사용합니다.",
        "문항·모범답안·채점 기준·난이도·출제 근거를 검토하고 수정한 뒤 최종 확정합니다.",
        "확정 후 학생별 시험지와 교사용 답안 PDF를 출력합니다. 수동 채점과 피드백을 저장한 뒤 결과를 공개합니다.",
    ])
    add_note(doc, "공개 전 확인", "시험 문항과 점수는 교사가 결과 공개를 실행하기 전까지 학생 화면에 보이지 않습니다. 학생별 답안과 결과가 섞이지 않았는지 표본을 확인하세요.", "warning")

    add_chapter(doc, "5. 자기·동료평가", "첫 운영은 성적 자동 반영이 아니라 학생 피드백과 교사 참고자료로 사용합니다.")
    add_step(doc, steps, "평가 회차 만들기·열기", "핵심 행동 문항 4개와 선택 문항 최대 1개를 확인합니다.")
    add_bullets(doc, bullets, [
        "학생은 관찰 가능한 행동 기준 4단계로 답합니다.",
        "자기평가에는 해당 활동 기회 없음, 동료평가에는 판단하기 어려움과 사유가 있습니다.",
        "1·2단계 동료평가는 학생에게 공개되지 않는 교사용 관찰 근거를 함께 받습니다.",
    ])
    add_step(doc, steps, "입력 마감·익명 의견 검토", "모든 공개 의견을 승인하거나 숨긴 뒤 결과를 공개합니다.")
    add_screenshot(doc, "05-teacher-evaluation.png", "그림 5. 교사의 익명 의견 검토 화면", "익명 평가 의견의 검토 필요 표시, 최소 가림 승인과 숨김 버튼이 있는 교사 화면")
    add_bullets(doc, bullets, [
        "교사는 평가자·대상자·숫자 응답·원문·교사용 근거를 모두 볼 수 있습니다.",
        "학생에게는 평가자 이름을 공개하지 않습니다. 개인정보·욕설만 최소한으로 가리고 의미를 바꾸지 않습니다.",
        "항목별 유효 숫자 평가가 3건 이상일 때만 평균과 승인 의견을 공개합니다.",
        "기준이 부족한 학생은 개별 의견 대신 교사 종합 피드백을 작성해야 합니다.",
    ])
    add_note(doc, "결과 공개 기준", "검토되지 않은 공개 의견이 남아 있으면 결과를 공개할 수 없습니다. 공개 뒤에는 입력을 다시 열 수 없으므로 미완료 학생과 종합 피드백을 마지막으로 확인합니다.", "danger")

    add_chapter(doc, "6. 도움말·문제 해결·보안", "대시보드의 사용법 문의는 공식 안내만 검색하며 외부 AI를 호출하지 않습니다.")
    add_screenshot(doc, "06-teacher-help.png", "그림 6. 교사용 공식 사용법 문의", "준비물 모바일 링크 질문에 공식 준비물 안내를 근거로 답한 교사용 도움말 화면")
    add_bullets(doc, bullets, [
        "기능 이름을 포함해 질문하면 관련 공식 절차를 찾기 쉽습니다.",
        "도움말은 학생 개인정보·비밀번호·학생 작성 자료를 조회·전달·변경·삭제하지 않습니다.",
        "학생 자료 처리는 교사가 해당 관리 화면에서 대상과 권한을 직접 확인한 뒤 실행합니다.",
    ])
    add_status_table(doc, [
        ("학생 로그인 실패", "학번·계정 상태를 확인하고 필요할 때만 비밀번호를 초기화합니다."),
        ("팀 화면 없음", "현재 팀 배정과 계정 상태를 확인합니다."),
        ("AI 입력 잠김", "응답 생성이 끝날 때까지 기다립니다."),
        ("저장·전송 실패", "작성 내용을 지우지 말고 WiFi 복구 뒤 다시 시도합니다."),
        ("시트 전송 실패", "기존 행을 확인한 뒤 교사 화면에서 재전송합니다."),
        ("복원 버튼 없음", "현재 팀장 또는 교사 권한인지 확인합니다."),
        ("평가 공개 불가", "미검토 의견과 유효 평가 수, 필수 종합 피드백을 확인합니다."),
    ], headings=("문제", "확인 순서"))
    add_checklist(doc, bullets, "수업 종료 전 체크리스트", [
        "교사 확인 필요 필터에 남은 항목을 확인했다.", "학생이 작성 내용을 저장·제출했는지 확인했다.",
        "Google Sheet에 시험 데이터나 중복 행을 만들지 않았다.", "임시 비밀번호를 문서·메신저·AI에 장기 보관하지 않았다.",
        "실제 학생 자료가 화면 캡처나 배부 문서에 노출되지 않았다.",
    ])
    doc.save(path)


def build_student(path):
    chapters = [
        "1. 첫 로그인과 팀 확인", "2. AI와 탐구 방향 찾기", "3. 계획서와 준비물 신청",
        "4. 개인 일지와 팀 보고서", "5. 시험 결과", "6. 자기·동료평가와 문제 해결",
    ]
    doc, bullets, steps = base_document("학생용")
    add_cover(doc, "학생용", "첫 로그인부터 자기·동료평가까지 팀 탐구를 진행하는 방법", chapters)

    add_chapter(doc, "1. 첫 로그인과 팀 확인", "선생님에게 받은 로그인 카드로 접속하고, 본인만 아는 새 비밀번호를 만든 뒤 팀 정보를 확인합니다.")
    add_screenshot(doc, "00-login.png", "그림 1. 과탐실 AI 탐구 플랫폼 로그인 화면", "플랫폼 소개와 학번 또는 교사 아이디, 비밀번호 입력란이 있는 로그인 화면")
    add_step(doc, steps, "로그인", "5자리 학번과 임시 비밀번호를 입력합니다.")
    add_bullets(doc, bullets, [
        "처음 로그인하면 8자 이상이며 글자와 숫자가 포함된 새 비밀번호로 바꿉니다.",
        "비밀번호는 친구에게 알려 주거나 공용 태블릿에 저장하지 않습니다.",
        "비밀번호를 잊으면 선생님께 초기화를 요청합니다. 여러 번 추측하지 않습니다.",
    ])
    add_step(doc, steps, "팀 확인", "화면 상단의 반·팀명·팀원·팀장을 확인합니다.")
    add_bullets(doc, bullets, [
        "팀이 다르거나 팀이 없다고 나오면 선생님께 바로 알립니다.",
        "팀원 제거·이동은 선생님만 처리합니다. 다른 학생 계정으로 로그인하지 않습니다.",
    ])

    add_chapter(doc, "2. AI와 탐구 방향 찾기", "AI는 답을 대신 쓰는 도구가 아니라 우리 팀이 더 좋은 질문과 계획을 만들도록 돕는 탐구 도우미입니다.")
    add_screenshot(doc, "01-student-overview.png", "그림 2. 학생 탐구 화면의 팀 정보·탭·AI 대화", "익명 체험팀의 팀 배너와 이론 탐구, 계획, 준비물, 일지, 보고서, 시험, 평가 탭 및 AI 대화")
    add_step(doc, steps, "팀 관심사 입력", "궁금한 현상과 비교하고 싶은 조건을 팀의 말로 적습니다.")
    add_bullets(doc, bullets, [
        "좋은 시작 예: 어떤 조건을 바꾸고 무엇을 측정할지 드러나는 질문",
        "피해야 할 시작: 정답만 요구하거나 완성된 계획서 전체를 써 달라는 요청",
        "실명·학번·연락처·비밀번호를 AI 대화에 입력하지 않습니다.",
    ])
    add_step(doc, steps, "방향 3개 비교·선택", "연구 질문, 독립·종속·통제변인, 안전 요소를 팀원과 비교합니다.")
    add_bullets(doc, bullets, [
        "선택한 방향을 바탕으로 개념·측정 방법·오차·안전·실제 출처를 질문합니다.",
        "AI가 보여 준 출처는 직접 열어 제목·작성 주체·내용을 확인합니다.",
        "AI 답변 생성 중에는 입력이 잠시 잠깁니다. 완료된 뒤 다음 질문을 보냅니다.",
    ])
    add_note(doc, "AI 사용 원칙", "우리 팀의 생각을 먼저 적고, AI의 제안은 사실 확인과 팀 토의를 거쳐 사용합니다. 존재하지 않는 논문·사이트·실험 결과를 만들지 않습니다.", "info")

    add_chapter(doc, "3. 계획서와 준비물 신청", "팀원이 항목을 나누어 작성하고, 서로 확인한 뒤 선생님께 제출합니다.")
    add_step(doc, steps, "계획서 공동 작성", "같은 항목에는 한 번에 한 명만 입력합니다.")
    add_bullets(doc, bullets, [
        "연구 동기·목적, 이론적 배경과 실제 출처, 변인·방법, 일정·장소, 안전, 기대효과, 참고문헌을 작성합니다.",
        "다른 팀원이 편집 중인 항목에는 작성자 표시와 잠금이 나타납니다. 잠금이 풀린 뒤 수정합니다.",
        "항목을 저장한 뒤 팀원과 전체 내용을 확인하고 선생님께 제출합니다.",
    ])
    add_status_table(doc, [
        ("선생님 확인 중", "제출 완료. 검토 결과를 기다립니다."),
        ("수정 요청", "교사 피드백을 읽고 보완한 뒤 다시 제출합니다."),
        ("승인", "계획이 승인되어 일지·보고서 단계로 진행합니다."),
        ("재승인 필요", "승인 뒤 내용을 바꾸었습니다. 반드시 다시 제출합니다."),
    ])
    add_step(doc, steps, "준비물 신청", "품명·규격·단가·개수·배송비·상품 링크를 확인합니다.")
    add_bullets(doc, bullets, [
        "쇼핑 앱에서 복사한 공유 문구 전체를 상품 링크 칸에 붙여 넣어도 됩니다.",
        "지원 쇼핑몰 주소는 PC용 주소로 자동 정리됩니다. 오류가 나오면 PC 브라우저에서 상품 주소를 다시 복사합니다.",
        "조별 합계가 5만원을 넘으면 경고가 표시됩니다. 제출 전에 선생님과 필요성을 확인합니다.",
        "제출 후 시트 반영 상태를 확인하고, 전송 실패가 보이면 반복 제출하지 말고 선생님께 알립니다.",
    ])

    add_chapter(doc, "4. 개인 일지와 팀 보고서", "일지는 각자 기록하고, 보고서는 팀이 역할을 나누어 함께 완성합니다.")
    add_step(doc, steps, "개인 실험 일지", "계획 승인 후 실험 일지 탭이 열립니다.")
    add_bullets(doc, bullets, [
        "차시와 날짜, 오늘 한 일, 관찰 결과, 느낀 점과 궁금한 점을 본인이 작성합니다.",
        "사진은 차시당 최대 5장입니다. 사람 얼굴·학번·연락처가 나오지 않게 촬영합니다.",
        "일지는 본인과 선생님만 볼 수 있습니다. 같은 팀 친구의 개인 일지는 볼 수 없습니다.",
        "WiFi가 끊겨도 내용을 지우지 않습니다. 브라우저 임시 저장 안내를 확인하고 연결 뒤 다시 전송합니다.",
    ])
    add_step(doc, steps, "팀 최종보고서", "연구 목적부터 결론·참고문헌·부록까지 학교 양식 순서로 작성합니다.")
    add_bullets(doc, bullets, [
        "팀원별 역할을 적고 각자 맡은 항목을 작성합니다.",
        "같은 항목을 다른 팀원이 편집 중이면 잠금이 풀릴 때까지 기다립니다. 다른 항목은 동시에 작성할 수 있습니다.",
        "자료 정리·결과 분석·고찰에서 관찰 사실과 해석을 구분합니다.",
        "팀 전체가 확인한 뒤 제출하고, 수정 요청이 오면 보완해 다시 제출합니다.",
    ])
    add_step(doc, steps, "변경 이력과 복원", "저장하기 전 전체 상태가 이력으로 남습니다.")
    add_screenshot(doc, "03-document-history.png", "그림 3. 계획서·보고서 변경 이력", "익명 예시의 변경 이력 세 건과 복원 버튼")
    add_bullets(doc, bullets, [
        "이력은 팀원이 볼 수 있지만 복원은 현재 팀장과 선생님만 할 수 있습니다.",
        "복원하면 복원 직전 상태도 이력에 남고, 문서는 다시 제출·검토해야 합니다.",
        "일반 팀원은 복원 버튼이 보이지 않습니다. 팀장에게 이유와 원하는 시점을 먼저 설명합니다.",
    ])

    add_chapter(doc, "5. 시험 결과", "시험지는 종이로 응시하고, 선생님이 채점·피드백을 입력해 결과를 공개하면 학생 화면에서 확인합니다.", page_break=False)
    add_bullets(doc, bullets, [
        "시험 결과 공개 전에는 문항·점수·피드백이 학생 화면에 보이지 않습니다.",
        "공개 후 시험 결과 탭에서 본인의 점수와 문항별 피드백만 확인합니다.",
        "다른 학생의 문항이나 점수를 보려고 하지 않습니다. 결과에 질문이 있으면 선생님께 본인 결과를 보여 주며 문의합니다.",
        "개인화 문항은 본인의 일지와 보고서 역할을 바탕으로 만들어집니다. 다른 학생의 개인 일지는 사용되지 않습니다.",
    ])
    add_note(doc, "공정성 안내", "전체 공통·팀 공통·개인화 문항 수와 총점 구조는 학생별로 같게 유지됩니다. 특정 학생의 실명·학번은 AI 출제 입력에 전달되지 않습니다.", "info")

    add_chapter(doc, "6. 자기·동료평가와 문제 해결", "평가는 친구를 줄 세우는 활동이 아니라 실제 팀 활동을 돌아보고 다음 협력을 더 좋게 만드는 피드백입니다.")
    add_screenshot(doc, "04-student-evaluation.png", "그림 4. 행동 기준 4단계 자기평가", "팀 활동 참여 문항에 관찰 가능한 4단계와 해당 활동 기회 없음이 표시된 학생 자기평가 화면")
    add_step(doc, steps, "자기평가", "실제 행동과 가장 가까운 단계를 고릅니다.")
    add_bullets(doc, bullets, [
        "활동 기회가 없었다면 해당 활동 기회 없음을 선택하고 이유를 적습니다.",
        "기여한 점과 다음에 개선할 점을 구체적인 행동으로 작성합니다.",
    ])
    add_step(doc, steps, "동료평가", "같은 팀의 현재 팀원을 실제로 관찰한 행동만 평가합니다.")
    add_bullets(doc, bullets, [
        "판단하기 어렵다면 억지로 숫자를 선택하지 말고 판단하기 어려움과 이유를 적습니다.",
        "1·2단계를 선택하면 선생님만 보는 관찰 근거를 적습니다. 모욕·추측·소문은 쓰지 않습니다.",
        "익명 의견은 선생님이 승인하거나 최소한으로 가린 뒤에만 공개됩니다.",
        "항목별 유효 평가가 3건 미만이면 평균과 개별 의견 대신 선생님 종합 피드백을 받습니다.",
    ])
    add_status_table(doc, [
        ("로그인 안 됨", "학번과 비밀번호를 다시 확인하고, 계속 실패하면 선생님께 알립니다."),
        ("팀 화면 없음", "팀 배정을 선생님께 확인합니다."),
        ("입력 잠김", "AI 답변 또는 다른 팀원의 편집이 끝날 때까지 기다립니다."),
        ("저장 오류", "내용을 지우지 말고 WiFi를 확인한 뒤 다시 저장합니다."),
        ("모바일 링크 오류", "PC 브라우저에서 상품 페이지 주소를 다시 복사합니다."),
        ("복원 버튼 없음", "일반 팀원은 정상적으로 복원할 수 없습니다. 현재 팀장이나 선생님께 요청합니다."),
    ], headings=("상황", "할 일"))
    add_checklist(doc, bullets, "개인정보·안전 체크리스트", [
        "AI 대화에 실명·학번·연락처·비밀번호를 입력하지 않았다.",
        "실험 사진에 얼굴·학번표·연락처가 나오지 않는다.",
        "실제 확인한 출처만 계획서·보고서에 적었다.",
        "위험한 화학물질·불꽃·고전압·미생물·인체 적용은 선생님 확인을 받았다.",
        "친구의 비밀번호·개인 일지·평가 원문을 요구하거나 공유하지 않았다.",
    ])
    doc.save(path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_teacher(OUT / "2026-과탐실-교사용-상세-사용서.docx")
    build_student(OUT / "2026-과탐실-학생용-상세-사용서.docx")


if __name__ == "__main__":
    main()
